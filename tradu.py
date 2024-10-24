import streamlit as st
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
import io
import re
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert
import pdfplumber
import os
import tempfile

class EnhancedPDFTranslator:
    def __init__(self):
        self.supported_languages = {
            'Português': 'pt',
            'Inglês': 'en',
            'Espanhol': 'es',
            'Francês': 'fr',
            'Alemão': 'de',
            'Italiano': 'it'
        }

    def analyze_document(self, pdf_file):
        """Analisa a estrutura do documento para escolher o melhor método de tradução"""
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        structure = {
            'has_tables': False,
            'has_images': False,
            'has_complex_layout': False,
            'text_density': 0
        }
        
        for page in doc:
            # Verificar imagens
            if len(page.get_images()) > 0:
                structure['has_images'] = True
            
            # Verificar tabelas e layout complexo
            blocks = page.get_text("blocks")
            if len(blocks) > 10:  # Valor arbitrário para determinar complexidade
                structure['has_complex_layout'] = True
            
            # Análise básica de tabelas
            text = page.get_text()
            if text.count('|') > 5 or text.count('\t') > 5:  # Indicadores simples de tabela
                structure['has_tables'] = True
            
            # Densidade de texto
            structure['text_density'] += len(text)
        
        doc.close()
        return structure

    def translate_with_pdf2docx(self, pdf_file, target_lang):
        """Método de tradução usando conversão para DOCX"""
        with tempfile.TemporaryDirectory() as temp_dir:
            # Salvar PDF temporariamente
            temp_pdf = os.path.join(temp_dir, "temp.pdf")
            with open(temp_pdf, 'wb') as f:
                f.write(pdf_file.getvalue())

            # Converter para DOCX
            temp_docx = os.path.join(temp_dir, "temp.docx")
            converter = Converter(temp_pdf)
            converter.convert(temp_docx)
            converter.close()

            # Traduzir DOCX
            doc = Document(temp_docx)
            translator = GoogleTranslator(target=target_lang)

            # Traduzir parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    try:
                        translated_text = translator.translate(paragraph.text)
                        paragraph.text = translated_text
                    except Exception as e:
                        st.warning(f"Erro ao traduzir parágrafo: {str(e)}")

            # Traduzir tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            try:
                                translated_text = translator.translate(cell.text)
                                cell.text = translated_text
                            except Exception as e:
                                st.warning(f"Erro ao traduzir célula: {str(e)}")

            # Salvar DOCX traduzido
            translated_docx = os.path.join(temp_dir, "translated.docx")
            doc.save(translated_docx)

            # Converter de volta para PDF
            translated_pdf = os.path.join(temp_dir, "translated.pdf")
            convert(translated_docx, translated_pdf)

            # Ler o PDF traduzido
            with open(translated_pdf, 'rb') as f:
                return f.read()

    def translate_with_pymupdf(self, pdf_file, target_lang):
        """Método de tradução usando PyMuPDF"""
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        new_doc = fitz.open()
        translator = GoogleTranslator(target=target_lang)

        for page_num in range(len(doc)):
            page = doc[page_num]
            new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
            new_page.show_pdf_page(new_page.rect, doc, page_num)
            
            blocks = page.get_text("blocks")
            for block in blocks:
                if block[6] == 0:  # Bloco de texto
                    try:
                        rect = fitz.Rect(block[:4])
                        text = block[4]
                        translated_text = translator.translate(text)
                        
                        # Preservar formatação
                        font_size = 10  # Detectar tamanho da fonte original
                        font_name = "helv"  # Detectar fonte original
                        
                        new_page.draw_rect(rect, color=None, fill=(1, 1, 1))
                        new_page.insert_text(
                            rect.tl,
                            translated_text,
                            fontname=font_name,
                            fontsize=font_size,
                            color=(0, 0, 0)
                        )
                    except Exception as e:
                        st.warning(f"Erro ao traduzir bloco: {str(e)}")

        pdf_bytes = new_doc.tobytes()
        doc.close()
        new_doc.close()
        return pdf_bytes

    def translate_pdf(self, pdf_file, target_lang):
        """Método principal de tradução que escolhe a melhor abordagem"""
        # Fazer uma cópia do arquivo em memória
        pdf_copy = io.BytesIO(pdf_file.read())
        pdf_file.seek(0)  # Resetar o ponteiro do arquivo

        # Analisar estrutura do documento
        structure = self.analyze_document(pdf_copy)
        
        try:
            if structure['has_tables'] or structure['has_complex_layout']:
                return self.translate_with_pdf2docx(pdf_file, target_lang)
            else:
                return self.translate_with_pymupdf(pdf_file, target_lang)
        except Exception as e:
            st.error(f"Erro no método principal de tradução: {str(e)}")
            # Tentar método alternativo
            try:
                if structure['has_tables']:
                    return self.translate_with_pymupdf(pdf_file, target_lang)
                else:
                    return self.translate_with_pdf2docx(pdf_file, target_lang)
            except Exception as e:
                raise Exception(f"Todos os métodos de tradução falharam: {str(e)}")

    def clean_filename(self, filename):
        return re.sub(r'[<>:"/\\|?*]', '', filename)

def main():
    st.set_page_config(page_title="PDF Translator", layout="wide")
    
    st.title("📚 Tradutor de PDF Avançado")
    st.markdown("### Traduza seus documentos PDF mantendo a estrutura original")
    
    translator = EnhancedPDFTranslator()
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Upload do PDF")
        pdf_file = st.file_uploader("Escolha um arquivo PDF", type=['pdf'])
        
        target_language = st.selectbox(
            "Selecione o idioma de destino",
            list(translator.supported_languages.keys())
        )
        
        if st.button("Traduzir"):
            if pdf_file is not None:
                with st.spinner("Analisando e traduzindo o documento... Isso pode levar alguns minutos."):
                    try:
                        target_lang_code = translator.supported_languages[target_language]
                        translated_pdf = translator.translate_pdf(pdf_file, target_lang_code)
                        
                        st.success("Tradução concluída!")
                        
                        original_filename = pdf_file.name
                        translated_filename = f"traduzido_{original_filename}"
                        
                        st.download_button(
                            label="Download do PDF traduzido",
                            data=translated_pdf,
                            file_name=translated_filename,
                            mime="application/pdf"
                        )
                            
                    except Exception as e:
                        st.error(f"Erro ao processar o PDF: {str(e)}")
            else:
                st.warning("Por favor, faça upload de um arquivo PDF primeiro.")

    st.markdown("---")
    st.markdown("""
    ### Como usar:
    1. Faça upload do seu arquivo PDF
    2. Selecione o idioma de destino
    3. Clique em "Traduzir"
    4. Baixe o PDF traduzido
    
    ### Recursos Avançados:
    - Análise automática da estrutura do documento
    - Seleção inteligente do método de tradução
    - Preservação aprimorada de layout e formatação
    - Suporte a documentos complexos com tabelas e imagens
    
    ### Observações:
    - A estrutura original do PDF será mantida (formatação, imagens, layout)
    - O processo pode levar alguns minutos dependendo do tamanho e complexidade
    - Documentos muito complexos podem requerer ajustes manuais
    """)

if __name__ == "__main__":
    main()
